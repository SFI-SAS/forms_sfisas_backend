import csv
from enum import Enum
from io import BytesIO
import io
from fastapi import APIRouter, HTTPException, Depends, Query
from fastapi.responses import StreamingResponse
import pandas as pd
from sqlalchemy.orm import Session, joinedload, selectinload
from sqlalchemy import and_, func, or_, select
from typing import List, Optional, Dict, Any
from pydantic import BaseModel
from datetime import datetime
from docx import Document
from app.core.security import get_current_user
from app.database import get_db
from app.models import Answer, Form, FormAnswer, FormApproval, FormApprovalNotification, FormCloseConfig, FormModerators, FormQuestion, FormSchedule, Question, QuestionFilterCondition, QuestionLocationRelation, QuestionTableRelation, QuestionType, Response, ResponseApproval, User
from app.schemas import DownloadRequest,FilterCondition

from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4

router = APIRouter()


# Schemas de respuesta
class UserInfo(BaseModel):
    id: int
    name: str
    email: str
    user_type: str
    nickname: Optional[str] = None
    category_name: Optional[str] = None

class QuestionOptionInfo(BaseModel):
    id: int
    option_text: str

class QuestionInfo(BaseModel):
    id: int
    question_text: str
    question_type: str
    required: bool
    root: bool
    category_name: Optional[str] = None
    options: List[QuestionOptionInfo] = []

class AnswerInfo(BaseModel):
    id: int
    question_id: int
    answer_text: Optional[str] = None
    file_path: Optional[str] = None
    file_serial: Optional[str] = None

class ResponseInfo(BaseModel):
    id: int
    user_id: int
    user_name: str
    mode: str
    mode_sequence: int
    repeated_id: Optional[str] = None
    submitted_at: datetime
    answers: List[AnswerInfo] = []

class FormModeratorInfo(BaseModel):
    id: int
    user_id: int
    user_name: str
    assigned_at: datetime

class FormScheduleInfo(BaseModel):
    id: int
    user_id: int
    user_name: str
    frequency_type: str
    repeat_days: Optional[str] = None
    interval_days: Optional[int] = None
    specific_date: Optional[datetime] = None
    status: bool

class FormApprovalInfo(BaseModel):
    id: int
    user_id: int
    user_name: str
    sequence_number: int
    is_mandatory: bool
    deadline_days: Optional[int] = None
    is_active: bool

class ResponseApprovalInfo(BaseModel):
    id: int
    response_id: int
    user_id: int
    user_name: str
    sequence_number: int
    is_mandatory: bool
    status: str
    reviewed_at: Optional[datetime] = None
    message: Optional[str] = None

class FormNotificationInfo(BaseModel):
    id: int
    user_id: int
    user_name: str
    notify_on: str

class FormCloseConfigInfo(BaseModel):
    id: int
    send_download_link: bool
    send_pdf_attachment: bool
    generate_report: bool
    do_nothing: bool
    download_link_recipient: Optional[str] = None
    email_recipient: Optional[str] = None
    report_recipient: Optional[str] = None

class QuestionFilterConditionInfo(BaseModel):
    id: int
    filtered_question_id: int
    source_question_id: int
    condition_question_id: int
    expected_value: str
    operator: str

class QuestionLocationRelationInfo(BaseModel):
    id: int
    origin_question_id: int
    target_question_id: int

class QuestionTableRelationInfo(BaseModel):
    id: int
    question_id: int
    related_question_id: Optional[int] = None
    name_table: str
    field_name: Optional[str] = None

class FormCompleteInfo(BaseModel):
    # Información básica del formulario
    id: int
    title: str
    description: Optional[str] = None
    format_type: str
    created_at: datetime
    # CORREGIDO: Cambié de Dict[str, Any] a List[Dict[str, Any]]
    # form_design: Optional[List[Dict[str, Any]]] = None
    
    # Información del creador
    creator: UserInfo
    
    # Preguntas del formulario
    questions: List[QuestionInfo] = []
    
    # Moderadores
    moderators: List[FormModeratorInfo] = []
    
    # Programación
    schedules: List[FormScheduleInfo] = []
    
    # Configuración de aprobaciones
    approval_templates: List[FormApprovalInfo] = []
    
    # Notificaciones
    notifications: List[FormNotificationInfo] = []
    
    # Configuración de cierre
    close_config: Optional[FormCloseConfigInfo] = None
    
    # Respuestas
    responses: List[ResponseInfo] = []
    
    # Aprobaciones de respuestas
    response_approvals: List[ResponseApprovalInfo] = []
    
    # Condiciones de filtro
    filter_conditions: List[QuestionFilterConditionInfo] = []
    
    # Relaciones de ubicación
    location_relations: List[QuestionLocationRelationInfo] = []
    
    # Relaciones de tabla
    table_relations: List[QuestionTableRelationInfo] = []
    
    # Respuestas del formulario
    form_answers: List[Dict[str, Any]] = []

@router.get("/forms/{form_id}/complete-info", response_model=FormCompleteInfo)
async def get_form_complete_info(form_id: int, db: Session = Depends(get_db),current_user: User = Depends(get_current_user)):
    
    """
    Obtiene toda la información relacionada con un formulario específico
    """
    if current_user == None:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="User does not have permission to get form"
        )
    # Obtener el formulario con todas sus relaciones
    form = db.query(Form).options(
        # Cargar el usuario creador con su categoría
        joinedload(Form.user).joinedload(User.category),
        
        # Cargar preguntas con sus opciones y categorías
        selectinload(Form.questions).selectinload(Question.options),
        selectinload(Form.questions).selectinload(Question.category),
        
        # Cargar moderadores con información del usuario
        selectinload(Form.form_moderators).selectinload(FormModerators.user),
        
        # Cargar respuestas con usuarios y respuestas
        selectinload(Form.responses).selectinload(Response.user),
        selectinload(Form.responses).selectinload(Response.answers).selectinload(Answer.question),
        selectinload(Form.responses).selectinload(Response.answers).selectinload(Answer.file_serial),
        
        # Cargar form_answers
        selectinload(Form.form_answers).selectinload(FormAnswer.question),
        
    ).filter(Form.id == form_id).first()
    
    if not form:
        raise HTTPException(status_code=404, detail="Formulario no encontrado")
    
    # Obtener información adicional relacionada con consultas separadas para obtener usuarios
    schedules = db.query(FormSchedule).filter(FormSchedule.form_id == form_id).all()
    
    # Obtener usuarios para schedules
    schedule_user_ids = [sched.user_id for sched in schedules]
    schedule_users = {}
    if schedule_user_ids:
        users = db.query(User).filter(User.id.in_(schedule_user_ids)).all()
        schedule_users = {user.id: user for user in users}
    
    approval_templates = db.query(FormApproval).filter(FormApproval.form_id == form_id).all()
    
    # Obtener usuarios para approval templates
    approval_user_ids = [approval.user_id for approval in approval_templates]
    approval_users = {}
    if approval_user_ids:
        users = db.query(User).filter(User.id.in_(approval_user_ids)).all()
        approval_users = {user.id: user for user in users}
    
    notifications = db.query(FormApprovalNotification).filter(FormApprovalNotification.form_id == form_id).all()
    
    # Obtener usuarios para notifications
    notification_user_ids = [notif.user_id for notif in notifications]
    notification_users = {}
    if notification_user_ids:
        users = db.query(User).filter(User.id.in_(notification_user_ids)).all()
        notification_users = {user.id: user for user in users}
    
    close_config = db.query(FormCloseConfig).filter(
        FormCloseConfig.form_id == form_id
    ).first()
    
    # Obtener aprobaciones de respuestas
    response_approvals = db.query(ResponseApproval).join(Response).filter(Response.form_id == form_id).all()
    
    # Obtener usuarios para response approvals
    response_approval_user_ids = [approval.user_id for approval in response_approvals]
    response_approval_users = {}
    if response_approval_user_ids:
        users = db.query(User).filter(User.id.in_(response_approval_user_ids)).all()
        response_approval_users = {user.id: user for user in users}
    
    # Obtener condiciones de filtro
    filter_conditions = db.query(QuestionFilterCondition).filter(
        QuestionFilterCondition.form_id == form_id
    ).all()
    
    # Obtener relaciones de ubicación
    location_relations = db.query(QuestionLocationRelation).filter(
        QuestionLocationRelation.form_id == form_id
    ).all()
    
    # Obtener relaciones de tabla para las preguntas del formulario
    question_ids = [q.id for q in form.questions]
    table_relations = db.query(QuestionTableRelation).filter(
        QuestionTableRelation.question_id.in_(question_ids)
    ).all()
    
    # Construir la respuesta
    result = FormCompleteInfo(
        id=form.id,
        title=form.title,
        description=form.description,
        format_type=form.format_type.value,
        created_at=form.created_at,
        
        # Información del creador
        creator=UserInfo(
            id=form.user.id,
            name=form.user.name,
            email=form.user.email,
            user_type=form.user.user_type.value,
            nickname=form.user.nickname,
            category_name=form.user.category.name if form.user.category else None
        ),
        
        # Preguntas
        questions=[
            QuestionInfo(
                id=q.id,
                question_text=q.question_text,
                question_type=q.question_type.value,
                required=q.required,
                root=q.root,
                category_name=q.category.name if q.category else None,
                options=[
                    QuestionOptionInfo(
                        id=opt.id,
                        option_text=opt.option_text
                    ) for opt in q.options
                ]
            ) for q in form.questions
        ],
        
        # Moderadores
        moderators=[
            FormModeratorInfo(
                id=mod.id,
                user_id=mod.user.id,
                user_name=mod.user.name,
                assigned_at=mod.assigned_at
            ) for mod in form.form_moderators
        ],
        
        # Programación
        schedules=[
            FormScheduleInfo(
                id=sched.id,
                user_id=sched.user_id,
                user_name=schedule_users[sched.user_id].name if sched.user_id in schedule_users else "Usuario no encontrado",
                frequency_type=sched.frequency_type,
                repeat_days=sched.repeat_days,
                interval_days=sched.interval_days,
                specific_date=sched.specific_date,
                status=sched.status
            ) for sched in schedules
        ],
        
        # Configuración de aprobaciones
        approval_templates=[
            FormApprovalInfo(
                id=approval.id,
                user_id=approval.user_id,
                user_name=approval_users[approval.user_id].name if approval.user_id in approval_users else "Usuario no encontrado",
                sequence_number=approval.sequence_number,
                is_mandatory=approval.is_mandatory,
                deadline_days=approval.deadline_days,
                is_active=approval.is_active
            ) for approval in approval_templates
        ],
        
        # Notificaciones
        notifications=[
            FormNotificationInfo(
                id=notif.id,
                user_id=notif.user_id,
                user_name=notification_users[notif.user_id].name if notif.user_id in notification_users else "Usuario no encontrado",
                notify_on=notif.notify_on
            ) for notif in notifications
        ],
        
        # Configuración de cierre
        close_config=FormCloseConfigInfo(
            id=close_config.id,
            send_download_link=close_config.send_download_link,
            send_pdf_attachment=close_config.send_pdf_attachment,
            generate_report=close_config.generate_report,
            do_nothing=close_config.do_nothing,
            download_link_recipient=close_config.download_link_recipient,
            email_recipient=close_config.email_recipient,
            report_recipient=close_config.report_recipient
        ) if close_config else None,
        
        # Respuestas
        responses=[
            ResponseInfo(
                id=resp.id,
                user_id=resp.user.id,
                user_name=resp.user.name,
                mode=resp.mode,
                mode_sequence=resp.mode_sequence,
                repeated_id=resp.repeated_id,
                submitted_at=resp.submitted_at,
                answers=[
                    AnswerInfo(
                        id=ans.id,
                        question_id=ans.question.id,
                        answer_text=ans.answer_text,
                        file_path=ans.file_path,
                        file_serial=ans.file_serial.serial if ans.file_serial else None
                    ) for ans in resp.answers
                ]
            ) for resp in form.responses
        ],
        
        # Aprobaciones de respuestas
        response_approvals=[
            ResponseApprovalInfo(
                id=approval.id,
                response_id=approval.response_id,
                user_id=approval.user_id,
                user_name=response_approval_users[approval.user_id].name if approval.user_id in response_approval_users else "Usuario no encontrado",
                sequence_number=approval.sequence_number,
                is_mandatory=approval.is_mandatory,
                status=approval.status.value,
                reviewed_at=approval.reviewed_at,
                message=approval.message
            ) for approval in response_approvals
        ],
        
        # Condiciones de filtro
        filter_conditions=[
            QuestionFilterConditionInfo(
                id=cond.id,
                filtered_question_id=cond.filtered_question_id,
                source_question_id=cond.source_question_id,
                condition_question_id=cond.condition_question_id,
                expected_value=cond.expected_value,
                operator=cond.operator
            ) for cond in filter_conditions
        ],
        
        # Relaciones de ubicación
        location_relations=[
            QuestionLocationRelationInfo(
                id=rel.id,
                origin_question_id=rel.origin_question_id,
                target_question_id=rel.target_question_id
            ) for rel in location_relations
        ],
        
        # Relaciones de tabla
        table_relations=[
            QuestionTableRelationInfo(
                id=rel.id,
                question_id=rel.question_id,
                related_question_id=rel.related_question_id,
                name_table=rel.name_table,
                field_name=rel.field_name
            ) for rel in table_relations
        ],
        
        # Form answers
        form_answers=[
            {
                "id": fa.id,
                "question_id": fa.question.id,
                "question_text": fa.question.question_text,
                "is_repeated": fa.is_repeated
            } for fa in form.form_answers
        ]
    )
    
    return result




@router.get("/forms/available")
async def get_available_forms(db: Session = Depends(get_db)):
    """Obtiene todos los formularios disponibles para descarga"""
    forms = db.query(Form).options(
        joinedload(Form.category),
        joinedload(Form.questions)
    ).all()
    
    return [
        {
            "id": form.id,
            "title": form.title,
            "description": form.description,
            "category": form.category.name if form.category else None,
            "total_responses": len(form.responses),
            "created_at": form.created_at
        }
        for form in forms
    ]
    
@router.get("/forms/{form_id}/fields")
async def get_form_fields(form_id: int, db: Session = Depends(get_db)):
    """Obtiene todos los campos/preguntas de un formulario específico con sus labels desde form_design"""
    form = db.query(Form).filter(Form.id == form_id).first()
    if not form:
        raise HTTPException(status_code=404, detail="Formulario no encontrado")
    
    # Parsear form_design para crear un mapa de id_question -> label
    label_map = {}
    if form.form_design:
        # form_design ya está deserializado como dict/list gracias a AutoJSON
        design_elements = form.form_design if isinstance(form.form_design, list) else []
        
        for element in design_elements:
            # Verificar que sea un elemento válido con id_question
            if isinstance(element, dict) and "id_question" in element:
                question_id = element.get("id_question")
                # Extraer el label desde props
                props = element.get("props", {})
                label = props.get("label", "")
                
                if question_id and label:
                    label_map[question_id] = label
    
    fields = []
    for question in form.questions:
        # Usar el label desde form_design, o fallback al question_text original
        display_label = label_map.get(question.id, question.question_text)
        
        field_info = {
            "id": question.id,
            "text": display_label,  # Ahora usa el label desde form_design
            "type": question.question_type.value,
            "required": question.required,
            "category": question.category.name if question.category else None
        }
        
        # Si es multiple choice, agregar opciones
        if question.question_type in [QuestionType.multiple_choice, QuestionType.one_choice]:
            field_info["options"] = [
                {"id": opt.id, "text": opt.option_text} 
                for opt in question.options
            ]
        
        fields.append(field_info)
    
    return {
        "form_id": form_id,
        "form_title": form.title,
        "fields": fields,
        "date_range": {
            "min_date": db.query(func.min(Response.submitted_at))
                         .filter(Response.form_id == form_id).scalar(),
            "max_date": db.query(func.max(Response.submitted_at))
                         .filter(Response.form_id == form_id).scalar()
        }
    }
# REEMPLAZAR COMPLETAMENTE la función preview_download_data:

@router.post("/download/preview")
async def preview_download_data(
    request: DownloadRequest, 
    db: Session = Depends(get_db)
):
    """Genera una vista previa de los datos que se descargarán"""
    
    # Construir query base
    query = db.query(Response).filter(Response.form_id.in_(request.form_ids))
    
    # Aplicar filtro de fecha si existe
    if request.date_filter:
        if request.date_filter.start_date:
            query = query.filter(Response.submitted_at >= request.date_filter.start_date)
        if request.date_filter.end_date:
            query = query.filter(Response.submitted_at <= request.date_filter.end_date)
    
    # Aplicar condiciones personalizadas de forma inteligente
    query = apply_smart_conditions(query, request.conditions, request.form_ids, db)
    
    # Obtener respuestas
    responses = query.limit(request.limit).all()
    
    # Obtener SOLO las preguntas seleccionadas
    questions = db.query(Question).filter(Question.id.in_(request.selected_fields)).all()
    question_dict = {q.id: q.question_text for q in questions}
    
    # Formatear datos SOLO con campos seleccionados y en orden
    preview_data = []
    for response in responses:
        # Crear objeto ordenado comenzando con campos base
        row_data = []
        
        # 1. Fecha Envío (siempre primera)
        row_data.append(response.submitted_at.strftime("%Y-%m-%d %H:%M:%S"))
        
        # 2. Formulario (siempre segunda)  
        row_data.append(response.form.title)
        
        # 3. Solo las preguntas seleccionadas EN SU ORDEN
        for field_id in request.selected_fields:
            # Buscar la respuesta específica para este field_id
            field_value = "-"
            for answer in response.answers:
                if answer.question_id == field_id:
                    field_value = answer.answer_text or answer.file_path or "-"
                    break
            row_data.append(field_value)
        
        preview_data.append(row_data)
    
    # Crear headers en el orden correcto
    headers = ["Fecha Envío", "Formulario"]
    for field_id in request.selected_fields:
        headers.append(question_dict.get(field_id, f"Pregunta_{field_id}"))
    
    # Convertir row_data arrays a objects para mantener compatibilidad
    formatted_preview_data = []
    for row_array in preview_data:
        row_obj = {}
        for i, header in enumerate(headers):
            if i < len(row_array):
                if i == 0:
                    row_obj["submitted_at"] = row_array[i]
                elif i == 1:
                    row_obj["form_title"] = row_array[i]
                else:
                    # Para las preguntas, usar el field_id correspondiente
                    field_index = i - 2
                    if field_index < len(request.selected_fields):
                        field_id = request.selected_fields[field_index]
                        row_obj[f"question_{field_id}"] = row_array[i]
        formatted_preview_data.append(row_obj)
    
    return {
        "total_records": query.count(),
        "preview_records": len(formatted_preview_data),
        "data": formatted_preview_data,
        "columns": headers,
        "applied_conditions_summary": get_conditions_summary(request.conditions, request.form_ids, db)
    }
    
def get_conditions_summary(conditions: List[FilterCondition], form_ids: List[int], db: Session):
    """
    Retorna un resumen de qué condiciones se aplicaron a qué formularios
    """
    summary = []
    
    for condition in conditions:
        # Obtener nombre de la pregunta
        question = db.query(Question).filter(Question.id == condition.field_id).first()
        question_text = question.question_text if question else f"Pregunta ID {condition.field_id}"
        
        # Determinar formularios afectados
        if condition.target_form_ids:
            affected_forms = condition.target_form_ids
        else:
            affected_forms = db.query(FormQuestion.form_id).filter(
                FormQuestion.question_id == condition.field_id,
                FormQuestion.form_id.in_(form_ids)
            ).all()
            affected_forms = [f.form_id for f in affected_forms]
        
        # Obtener nombres de formularios
        if affected_forms:
            form_names = db.query(Form.title).filter(Form.id.in_(affected_forms)).all()
            form_names = [f.title for f in form_names]
        else:
            form_names = []
        
        summary.append({
            "condition": f"{question_text} {condition.operator} '{condition.value}'",
            "applied_to_forms": form_names,
            "affected_form_count": len(form_names)
        })
    
    return summary


@router.get("/forms/fields-analysis")
async def analyze_form_fields(
    form_ids: List[int] = Query(...), 
    db: Session = Depends(get_db)
):
    """
    Analiza qué campos están disponibles en qué formularios.
    Útil para el frontend para mostrar qué condiciones se pueden aplicar.
    """
    if not form_ids:
        # Si no se proporcionan form_ids, obtener todos
        all_forms = db.query(Form.id).all()
        form_ids = [f.id for f in all_forms]
    
    # Obtener todas las preguntas usadas en los formularios seleccionados
    questions_in_forms = db.query(
        FormQuestion.form_id,
        FormQuestion.question_id,
        Question.question_text,
        Question.question_type,
        Form.title
    ).join(
        Question, FormQuestion.question_id == Question.id
    ).join(
        Form, FormQuestion.form_id == Form.id
    ).filter(
        FormQuestion.form_id.in_(form_ids)
    ).all()
    
    # Organizar datos por pregunta
    fields_analysis = {}
    for fq in questions_in_forms:
        if fq.question_id not in fields_analysis:
            fields_analysis[fq.question_id] = {
                "question_id": fq.question_id,
                "question_text": fq.question_text,
                "question_type": fq.question_type.value,
                "available_in_forms": [],
                "form_count": 0
            }
        
        fields_analysis[fq.question_id]["available_in_forms"].append({
            "form_id": fq.form_id,
            "form_title": fq.title
        })
        fields_analysis[fq.question_id]["form_count"] += 1
    
    # Convertir a lista y ordenar por uso más común
    result = list(fields_analysis.values())
    result.sort(key=lambda x: x["form_count"], reverse=True)
    
    return {
        "total_forms_selected": len(form_ids),
        "fields_analysis": result,
        "common_fields": [f for f in result if f["form_count"] == len(form_ids)],
        "partial_fields": [f for f in result if f["form_count"] < len(form_ids)]
    }
def get_column_headers(field_ids: List[int], db: Session):
    """Obtiene los headers de las columnas EN EL ORDEN CORRECTO"""
    # Headers base fijos
    headers = ["Fecha Envío", "Formulario"]
    
    # Obtener preguntas y crear diccionario
    if field_ids:
        questions = db.query(Question).filter(Question.id.in_(field_ids)).all()
        question_dict = {q.id: q.question_text for q in questions}
        
        # Agregar headers EN EL ORDEN de field_ids
        for field_id in field_ids:
            if field_id in question_dict:
                headers.append(question_dict[field_id])
            else:
                headers.append(f"Pregunta_{field_id}")
    
    return headers

class DownloadFormat(str, Enum):
    excel = "excel"
    csv = "csv"
    pdf = "pdf"
    word = "word"

class FinalDownloadRequest(DownloadRequest):
    format: DownloadFormat

@router.post("/download/generate")
async def generate_download(
    request: FinalDownloadRequest,
    db: Session = Depends(get_db)
):
    """Genera el archivo de descarga en el formato solicitado"""
    
    # Reutilizar la lógica de preview pero sin límite
    data = await get_filtered_data(request, db, limit=None)
    
    if request.format == DownloadFormat.excel:
        return generate_excel_response(data)
    elif request.format == DownloadFormat.csv:
        return generate_csv_response(data)
    elif request.format == DownloadFormat.pdf:
        return generate_pdf_response(data)
    elif request.format == DownloadFormat.word:
        return generate_word_response(data)
# REEMPLAZAR la función get_filtered_data en tu archivo backend:

# REEMPLAZAR COMPLETAMENTE la función get_filtered_data:

async def get_filtered_data(request: FinalDownloadRequest, db: Session, limit: Optional[int] = None):
    """Función auxiliar que obtiene los datos filtrados con lógica inteligente"""
    # Construir query base
    query = db.query(Response).filter(Response.form_id.in_(request.form_ids))
    
    # Aplicar filtro de fecha si existe
    if request.date_filter:
        if request.date_filter.start_date:
            query = query.filter(Response.submitted_at >= request.date_filter.start_date)
        if request.date_filter.end_date:
            query = query.filter(Response.submitted_at <= request.date_filter.end_date)
    
    # Aplicar condiciones personalizadas de forma inteligente
    query = apply_smart_conditions(query, request.conditions, request.form_ids, db)
    
    # Aplicar límite si se especifica
    if limit:
        responses = query.limit(limit).all()
    else:
        responses = query.all()
    
    # Obtener SOLO las preguntas seleccionadas
    questions = db.query(Question).filter(Question.id.in_(request.selected_fields)).all()
    question_dict = {q.id: q.question_text for q in questions}
    
    # Formatear datos SOLO con campos seleccionados
    formatted_data = []
    for response in responses:
        row = {
            "submitted_at": response.submitted_at.strftime("%Y-%m-%d %H:%M:%S"),
            "form_title": response.form.title
        }
        
        # Agregar SOLO las preguntas seleccionadas EN SU ORDEN
        for field_id in request.selected_fields:
            column_name = question_dict.get(field_id, f"Pregunta_{field_id}")
            field_value = "-"
            for answer in response.answers:
                if answer.question_id == field_id:
                    field_value = answer.answer_text or answer.file_path or "-"
                    break
            row[column_name] = field_value
        
        formatted_data.append(row)
    
    # Crear lista de columnas en el orden correcto
    ordered_columns = ["submitted_at", "form_title"]
    for field_id in request.selected_fields:
        column_name = question_dict.get(field_id, f"Pregunta_{field_id}")
        ordered_columns.append(column_name)
    
    return {
        "data": formatted_data,
        "total_records": len(formatted_data),
        "columns": ordered_columns
    }
    

def apply_smart_conditions(query, conditions: List[FilterCondition], form_ids: List[int], db: Session):
    """
    Aplica condiciones de filtrado de forma inteligente.
    Solo aplica cada condición a los formularios que tienen el campo específico.
    """
    for condition in conditions:
        # Determinar a qué formularios aplicar esta condición
        if condition.target_form_ids:
            # Si se especificaron formularios específicos, usar esos
            target_forms = [fid for fid in condition.target_form_ids if fid in form_ids]
        else:
            # Si no se especificaron, encontrar automáticamente qué formularios tienen esta pregunta
            target_forms_query = db.query(FormQuestion.form_id).filter(
                FormQuestion.question_id == condition.field_id,
                FormQuestion.form_id.in_(form_ids)
            ).all()
            target_forms = [f.form_id for f in target_forms_query]
        
        if not target_forms:
            # Si ningún formulario tiene esta pregunta, saltar la condición
            continue
        
        # Crear subquery para respuestas que cumplen la condición
        subquery = db.query(Answer.response_id).filter(
            Answer.question_id == condition.field_id
        )
        
        # Aplicar el operador específico
        if condition.operator == "=":
            subquery = subquery.filter(Answer.answer_text == condition.value)
        elif condition.operator == "!=":
            subquery = subquery.filter(Answer.answer_text != condition.value)
        elif condition.operator == "contains":
            subquery = subquery.filter(Answer.answer_text.contains(condition.value))
        elif condition.operator == "starts_with":
            subquery = subquery.filter(Answer.answer_text.startswith(condition.value))
        elif condition.operator == "ends_with":
            subquery = subquery.filter(Answer.answer_text.endswith(condition.value))
        elif condition.operator == ">":
            subquery = subquery.filter(Answer.answer_text > condition.value)
        elif condition.operator == "<":
            subquery = subquery.filter(Answer.answer_text < condition.value)
        elif condition.operator == ">=":
            subquery = subquery.filter(Answer.answer_text >= condition.value)
        elif condition.operator == "<=":
            subquery = subquery.filter(Answer.answer_text <= condition.value)
        
        # Aplicar filtro inteligente:
        # - Si la respuesta es de un formulario que tiene la pregunta, debe cumplir la condición
        # - Si la respuesta es de un formulario que NO tiene la pregunta, se incluye sin filtro
        query = query.filter(
            or_(
                # Respuestas de formularios que tienen el campo y cumplen la condición
                and_(
                    Response.form_id.in_(target_forms),
                    Response.id.in_(subquery)
                ),
                # Respuestas de formularios que NO tienen el campo (se incluyen sin condición)
                ~Response.form_id.in_(target_forms)
            )
        )
    
    return query

def generate_excel_response(data: Dict):
    """Genera archivo Excel"""
    output = io.BytesIO()
    
    # Crear DataFrame
    df = pd.DataFrame(data['data'])
    
    # Escribir a Excel con formato
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df.to_excel(writer, sheet_name='Datos', index=False)
        
        # Obtener workbook y worksheet para aplicar formato
        workbook = writer.book
        worksheet = writer.sheets['Datos']
        
        # Formato para headers
        header_format = workbook.add_format({
            'bold': True,
            'text_wrap': True,
            'valign': 'top',
            'fg_color': '#D7E4BC',
            'border': 1
        })
        
        # Aplicar formato a headers
        for col_num, value in enumerate(df.columns.values):
            worksheet.write(0, col_num, value, header_format)
            worksheet.set_column(col_num, col_num, 20)  # Ancho de columna
        
        # Agregar filtros automáticos
        worksheet.autofilter(0, 0, len(df), len(df.columns) - 1)
    
    output.seek(0)
    
    return StreamingResponse(
        io.BytesIO(output.read()),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename=datos_formularios.xlsx"}
    )

def generate_csv_response(data: Dict):
    """Genera archivo CSV"""
    output = io.StringIO()
    df = pd.DataFrame(data['data'])
    df.to_csv(output, index=False, encoding='utf-8', sep=',')
    
    content = output.getvalue().encode('utf-8')
    
    return StreamingResponse(
        io.BytesIO(content),
        media_type="text/csv; charset=utf-8",
        headers={"Content-Disposition": "attachment; filename=datos_formularios.csv"}
    )

def generate_pdf_response(data: Dict):
    """Genera archivo PDF con campos optimizados y contenido bien ajustado"""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter, A4, landscape
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak, KeepTogether
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch, cm
    import math
    import textwrap
    
    output = io.BytesIO()
    
    # Usar landscape para máximo espacio
    pagesize = landscape(A4)
    page_width = landscape(A4)[0] - 3*cm
    page_height = landscape(A4)[1] - 3*cm
    
    doc = SimpleDocTemplate(
        output, 
        pagesize=pagesize,
        leftMargin=1.5*cm,
        rightMargin=1.5*cm,
        topMargin=2*cm,
        bottomMargin=2*cm
    )
    
    styles = getSampleStyleSheet()
    story = []
    
    # Título principal
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=16,
        textColor=colors.darkblue,
        spaceAfter=20,
        alignment=1,
        spaceBefore=0
    )
    story.append(Paragraph("Reporte Completo de Datos de Formularios", title_style))
    
    # Información del reporte
    info_style = ParagraphStyle(
        'InfoStyle',
        parent=styles['Normal'],
        fontSize=10,
        alignment=1,
        spaceAfter=15
    )
    info_text = f"Generado el {datetime.now().strftime('%d/%m/%Y a las %H:%M:%S')}<br/>"
    info_text += f"Total de registros: <b>{data['total_records']}</b>"
    if data['data']:
        info_text += f" | Total de campos: <b>{len(data['data'][0].keys())}</b>"
    
    story.append(Paragraph(info_text, info_style))
    story.append(Spacer(1, 15))
    
    if data['data']:
        headers = list(data['data'][0].keys())
        total_fields = len(headers)
        
        # Función mejorada para analizar contenido y ajustar texto
        def analyze_content_length(data, headers):
            """Analiza el contenido para determinar anchos óptimos de columna"""
            col_max_lengths = {}
            
            for header in headers:
                # Longitud del header
                header_length = len(str(header))
                max_content_length = 0
                
                # Analizar contenido de cada registro
                for record in data:
                    content = str(record.get(header, ''))
                    # Considerar saltos de línea
                    lines = content.split('\n')
                    max_line_length = max([len(line) for line in lines]) if lines else 0
                    max_content_length = max(max_content_length, max_line_length)
                
                # El ancho mínimo será el mayor entre header y contenido, con límites
                col_max_lengths[header] = max(header_length, min(max_content_length, 50))
            
            return col_max_lengths
        
        def smart_text_wrap(text, max_width=None, max_lines=3):
            """Envuelve texto inteligentemente preservando información importante"""
            if text is None:
                return ''
            
            text = str(text).strip()
            if not text:
                return ''
            
            # Reemplazar saltos de línea múltiples por espacios
            text = ' '.join(text.split())
            
            if max_width is None:
                max_width = 40
                
            # Si el texto es corto, devolverlo tal como está
            if len(text) <= max_width:
                return text
            
            # Usar textwrap para dividir en líneas
            lines = textwrap.wrap(text, width=max_width, max_lines=max_lines)
            
            if len(lines) <= max_lines:
                return '\n'.join(lines)
            else:
                # Si se excede las líneas, truncar la última línea
                result_lines = lines[:max_lines-1]
                last_line = lines[max_lines-1]
                if len(last_line) > max_width - 3:
                    last_line = last_line[:max_width-3] + '...'
                else:
                    last_line = last_line + '...'
                result_lines.append(last_line)
                return '\n'.join(result_lines)
        
        def format_header(header, max_width=20):
            """Formatea headers para que sean legibles"""
            if len(header) <= max_width:
                return header
            
            # Intentar dividir por separadores comunes
            for separator in ['_', '-', ' ', '.']:
                if separator in header:
                    parts = header.split(separator)
                    if len(parts) > 1:
                        # Tomar las primeras palabras que quepan
                        result = parts[0]
                        for part in parts[1:]:
                            if len(result + separator + part) <= max_width - 3:
                                result += separator + part
                            else:
                                result += '...'
                                break
                        return result
            
            # Si no hay separadores, truncar directamente
            return header[:max_width-3] + '...'
        
        # Analizar contenido para optimizar anchos
        col_lengths = analyze_content_length(data['data'], headers)
        
        # Configuración mejorada de columnas
        MIN_COL_WIDTH = 2.5 * cm  # Ancho mínimo aumentado
        OPTIMAL_COL_WIDTH = 4 * cm  # Ancho óptimo
        MAX_COL_WIDTH = 6 * cm  # Ancho máximo
        
        # Calcular anchos de columna basados en contenido
        def calculate_column_widths(headers, col_lengths, available_width):
            """Calcula anchos óptimos de columna basados en contenido"""
            total_chars = sum(col_lengths.values())
            widths = {}
            
            for header in headers:
                # Proporción basada en contenido
                content_ratio = col_lengths[header] / total_chars if total_chars > 0 else 1/len(headers)
                proposed_width = available_width * content_ratio
                
                # Aplicar límites
                if proposed_width < MIN_COL_WIDTH:
                    widths[header] = MIN_COL_WIDTH
                elif proposed_width > MAX_COL_WIDTH:
                    widths[header] = MAX_COL_WIDTH
                else:
                    widths[header] = proposed_width
            
            # Ajustar si el total excede el ancho disponible
            total_width = sum(widths.values())
            if total_width > available_width:
                scale_factor = available_width / total_width
                for header in headers:
                    widths[header] *= scale_factor
                    # Respetar ancho mínimo después del escalado
                    widths[header] = max(widths[header], MIN_COL_WIDTH * 0.8)
            
            return [widths[header] for header in headers]
        
        # Calcular número máximo de columnas que caben
        min_total_width = len(headers) * MIN_COL_WIDTH
        
        if min_total_width > page_width:
            # ESTRATEGIA 1: Dividir por campos (demasiadas columnas)
            chars_per_page = page_width / (MIN_COL_WIDTH / 15)  # Aproximación de caracteres por página
            
            # Agrupar campos por longitud de contenido
            field_chunks = []
            current_chunk = []
            current_width = 0
            
            for header in headers:
                expected_width = min(col_lengths[header] * 0.15 * cm, MAX_COL_WIDTH)
                expected_width = max(expected_width, MIN_COL_WIDTH)
                
                if current_width + expected_width <= page_width and current_chunk:
                    current_chunk.append(header)
                    current_width += expected_width
                else:
                    if current_chunk:
                        field_chunks.append(current_chunk)
                    current_chunk = [header]
                    current_width = expected_width
            
            if current_chunk:
                field_chunks.append(current_chunk)
            
            # Procesar cada chunk de campos
            chunk_number = 1
            for field_chunk in field_chunks:
                section_style = ParagraphStyle(
                    'SectionTitle',
                    parent=styles['Heading2'],
                    fontSize=12,
                    textColor=colors.darkgreen,
                    spaceAfter=10,
                    spaceBefore=10
                )
                
                section_title = f"Sección {chunk_number}/{len(field_chunks)} - "
                section_title += f"Campos {headers.index(field_chunk[0])+1} al {headers.index(field_chunk[-1])+1}"
                section_title += f" ({len(field_chunk)} columnas)"
                
                story.append(Paragraph(section_title, section_style))
                
                # Calcular anchos optimizados para este chunk
                chunk_col_lengths = {h: col_lengths[h] for h in field_chunk}
                col_widths = calculate_column_widths(field_chunk, chunk_col_lengths, page_width)
                
                # Crear tabla para este chunk
                table_data = []
                
                # Headers
                header_row = []
                for header in field_chunk:
                    formatted_header = format_header(header, 25)
                    header_row.append(formatted_header)
                table_data.append(header_row)
                
                # Datos
                for record in data['data']:
                    row_data = []
                    for i, header in enumerate(field_chunk):
                        value = record.get(header, '')
                        # Calcular ancho máximo de caracteres para esta columna
                        char_width = int(col_widths[i] / (0.15 * cm))
                        wrapped_value = smart_text_wrap(value, char_width, 4)
                        row_data.append(wrapped_value)
                    table_data.append(row_data)
                
                # Crear y estilizar tabla
                table = Table(table_data, colWidths=col_widths, repeatRows=1)
                
                table.setStyle(TableStyle([
                    # Headers
                    ('BACKGROUND', (0, 0), (-1, 0), colors.darkblue),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 8),
                    ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                    
                    # Data rows
                    ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                    ('FONTSIZE', (0, 1), (-1, -1), 7),
                    ('ALIGN', (0, 1), (-1, -1), 'LEFT'),
                    
                    # Padding optimizado
                    ('TOPPADDING', (0, 0), (-1, -1), 6),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                    ('LEFTPADDING', (0, 0), (-1, -1), 4),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 4),
                    
                    # Bordes y colores
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey]),
                    
                    # CRÍTICO: Permitir que el texto se ajuste dentro de las celdas
                    ('WORDWRAP', (0, 0), (-1, -1), 'CJK'),
                    ('LEADING', (0, 1), (-1, -1), 10),  # Espaciado entre líneas
                ]))
                
                story.append(KeepTogether(table))
                
                if chunk_number < len(field_chunks):
                    story.append(PageBreak())
                
                chunk_number += 1
        
        else:
            # ESTRATEGIA 2: Todas las columnas caben en una página
            col_widths = calculate_column_widths(headers, col_lengths, page_width)
            
            # Preparar datos de la tabla
            table_data = []
            
            # Headers
            header_row = []
            for header in headers:
                formatted_header = format_header(header, 20)
                header_row.append(formatted_header)
            table_data.append(header_row)
            
            # Datos con ajuste inteligente
            for record in data['data']:
                row_data = []
                for i, header in enumerate(headers):
                    value = record.get(header, '')
                    # Calcular ancho de caracteres disponible
                    char_width = int(col_widths[i] / (0.15 * cm))
                    wrapped_value = smart_text_wrap(value, char_width, 3)
                    row_data.append(wrapped_value)
                table_data.append(row_data)
            
            # Crear tabla
            table = Table(table_data, colWidths=col_widths)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.darkblue),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 8),
                ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 1), (-1, -1), 7),
                ('ALIGN', (0, 1), (-1, -1), 'LEFT'),
                ('TOPPADDING', (0, 0), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                ('LEFTPADDING', (0, 0), (-1, -1), 4),
                ('RIGHTPADDING', (0, 0), (-1, -1), 4),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey]),
                ('WORDWRAP', (0, 0), (-1, -1), 'CJK'),
                ('LEADING', (0, 1), (-1, -1), 10),
            ]))
            
            story.append(table)
        
    
    else:
        story.append(Paragraph("⚠️ No hay datos disponibles para mostrar.", styles['Normal']))
    
    # Generar PDF
    doc.build(story)
    output.seek(0)
    
    return StreamingResponse(
        io.BytesIO(output.read()),
        media_type="application/pdf",
        headers={"Content-Disposition": "attachment; filename=reporte_optimizado_formularios.pdf"}
    )

def generate_word_response(data: Dict):
    """
    Genera archivo Word profesional con formato mejorado y manejo inteligente de datos
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from datetime import datetime
    import io
    
    # Crear documento Word
    doc = Document()
    
    # Configurar márgenes del documento
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Título principal
    title = doc.add_heading('REPORTE CONSOLIDADO DE FORMULARIOS', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = RGBColor(47, 85, 151)  # Azul profesional
    
    # Subtítulo
    subtitle = doc.add_heading('Vista Integral de Datos', level=2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.color.rgb = RGBColor(102, 102, 102)
    
    # Línea separadora
    doc.add_paragraph('_' * 80).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Información del reporte
    current_date = datetime.now()
    info_paragraph = doc.add_paragraph()
    info_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Formatear información con estilos
    info_paragraph.add_run('📊 INFORMACIÓN DEL REPORTE\n').bold = True
    info_paragraph.add_run(f'Fecha de generación: ').bold = True
    info_paragraph.add_run(f'{current_date.strftime("%d/%m/%Y %H:%M:%S")}\n')
    info_paragraph.add_run(f'Total de registros: ').bold = True
    info_paragraph.add_run(f'{data.get("total_records", len(data.get("data", []))):,}\n')
    
    if data.get('data') and len(data['data']) > 0:
        # Obtener información de columnas
        headers = list(data['data'][0].keys())
        num_columns = len(headers)
        
        info_paragraph.add_run(f'Campos disponibles: ').bold = True
        info_paragraph.add_run(f'{num_columns}\n')
        info_paragraph.add_run(f'Formato: ').bold = True
        
        # Determinar estrategia según número de columnas
        if num_columns <= 8:
            strategy = "Tabla única con todas las columnas"
            max_records_show = 50
        elif num_columns <= 15:
            strategy = "Tabla única con formato compacto"
            max_records_show = 30
        else:
            strategy = "Tablas divididas por secciones"
            max_records_show = 25
        
        info_paragraph.add_run(f'{strategy}\n')
        
        # Espacio
        doc.add_paragraph()
        
        # Limpiar nombres de headers
        def clean_header_name(header):
            return str(header).replace('_', ' ').title()
        
        clean_headers = [clean_header_name(h) for h in headers]
        
        if num_columns <= 15:
            # TABLA ÚNICA
            doc.add_heading('📋 DATOS CONSOLIDADOS', level=2)
            
            # Crear tabla
            table = doc.add_table(rows=1, cols=num_columns)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Headers
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(clean_headers):
                hdr_cells[i].text = header
                
                # Estilo del header
                for paragraph in hdr_cells[i].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(10)
                        run.font.color.rgb = RGBColor(255, 255, 255)
                
                # Color de fondo del header
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), '2F5597')  # Azul
                hdr_cells[i]._tc.get_or_add_tcPr().append(shading_elm)
            
            # Datos (limitar registros para Word)
            records_to_show = min(len(data['data']), max_records_show)
            
            for idx, row_data in enumerate(data['data'][:records_to_show]):
                row_cells = table.add_row().cells
                for i, header in enumerate(headers):
                    cell_value = row_data.get(header, '')
                    
                    # Formatear valor
                    if not cell_value or str(cell_value).strip() == '':
                        display_value = '-'
                    else:
                        display_value = str(cell_value)
                        
                        # Formateo especial para algunos campos
                        if 'precio' in header.lower():
                            try:
                                if display_value.replace(',', '').isdigit():
                                    display_value = f"${int(display_value):,}"
                            except:
                                pass
                    
                    row_cells[i].text = display_value
                    
                    # Estilo de celda de datos
                    for paragraph in row_cells[i].paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        for run in paragraph.runs:
                            run.font.size = Pt(9)
                    
                    # Color alternado para filas
                    if idx % 2 == 1:
                        shading_elm = OxmlElement('w:shd')
                        shading_elm.set(qn('w:fill'), 'F8F9FA')  # Gris muy claro
                        row_cells[i]._tc.get_or_add_tcPr().append(shading_elm)
            
            # Ajustar ancho de columnas
            for i, header in enumerate(headers):
                # Ancho basado en contenido del header
                header_len = len(clean_headers[i])
                if header_len <= 8:
                    width = Inches(1.0)
                elif header_len <= 15:
                    width = Inches(1.5)
                else:
                    width = Inches(2.0)
                
                table.columns[i].width = width
            
            # Nota si hay más registros
            if len(data['data']) > records_to_show:
                doc.add_paragraph()
                note_paragraph = doc.add_paragraph()
                note_paragraph.add_run('📝 Nota: ').bold = True
                note_paragraph.add_run(f'Se muestran los primeros {records_to_show:,} registros de {len(data["data"]):,} totales. ')
                note_paragraph.add_run('Para ver todos los datos, utilice la exportación a Excel o CSV.')
        
        else:
            # TABLAS DIVIDIDAS - Para más de 15 columnas
            doc.add_heading('📋 DATOS DIVIDIDOS EN SECCIONES', level=2)
            
            # Dividir columnas en grupos
            group_size = 6
            
            # Identificar columnas clave
            key_headers = []
            other_headers = []
            
            for header in headers:
                if any(word in header.lower() for word in ['fecha', 'id', 'numero', 'formato']):
                    key_headers.append(header)
                else:
                    other_headers.append(header)
            
            # Crear grupos
            header_groups = []
            base_group = key_headers[:2]  # Máximo 2 identificadores
            
            remaining_space = group_size - len(base_group)
            for i in range(0, len(other_headers), remaining_space):
                group = base_group + other_headers[i:i + remaining_space]
                header_groups.append(group)
            
            # Crear tabla para cada grupo
            for group_idx, header_group in enumerate(header_groups):
                if group_idx > 0:
                    doc.add_page_break()
                
                # Título de la sección
                section_title = f"SECCIÓN {group_idx + 1}: {', '.join([clean_header_name(h) for h in header_group])}"
                doc.add_heading(section_title, level=3)
                
                clean_group_headers = [clean_header_name(h) for h in header_group]
                
                # Crear tabla para este grupo
                group_table = doc.add_table(rows=1, cols=len(header_group))
                group_table.style = 'Table Grid'
                group_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                # Headers del grupo
                hdr_cells = group_table.rows[0].cells
                for i, header in enumerate(clean_group_headers):
                    hdr_cells[i].text = header
                    
                    for paragraph in hdr_cells[i].paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = Pt(10)
                            run.font.color.rgb = RGBColor(255, 255, 255)
                    
                    # Color verde para diferenciar secciones
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), '1A472A')  # Verde oscuro
                    hdr_cells[i]._tc.get_or_add_tcPr().append(shading_elm)
                
                # Datos del grupo
                for idx, row_data in enumerate(data['data'][:max_records_show]):
                    row_cells = group_table.add_row().cells
                    for i, header in enumerate(header_group):
                        cell_value = row_data.get(header, '')
                        display_value = str(cell_value) if cell_value and str(cell_value).strip() else '-'
                        
                        row_cells[i].text = display_value
                        
                        for paragraph in row_cells[i].paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            for run in paragraph.runs:
                                run.font.size = Pt(9)
                        
                        # Color alternado
                        if idx % 2 == 1:
                            shading_elm = OxmlElement('w:shd')
                            shading_elm.set(qn('w:fill'), 'F0FDF4')  # Verde muy claro
                            row_cells[i]._tc.get_or_add_tcPr().append(shading_elm)
                
                # Ajustar anchos
                for i, header in enumerate(clean_group_headers):
                    group_table.columns[i].width = Inches(1.8)
                
                doc.add_paragraph()
        
    else:
        # Sin datos
        doc.add_paragraph()
        no_data_paragraph = doc.add_paragraph()
        no_data_paragraph.add_run('❌ No se encontraron datos para generar el reporte.').bold = True
        no_data_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Footer
    doc.add_paragraph()
    footer_paragraph = doc.add_paragraph()
    footer_paragraph.add_run('📄 Documento generado por Sistema Empresarial - Reporte automático optimizado')
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer_paragraph.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(128, 128, 128)
    
    # CORRECCIÓN PRINCIPAL: Manejo correcto del BytesIO
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    
    # Leer el contenido del buffer
    content = output.read()
    output.close()
    
    # Crear un nuevo BytesIO con el contenido
    final_output = io.BytesIO(content)
    
    return StreamingResponse(
        final_output,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": "attachment; filename=reporte_consolidado.docx",
            "Content-Type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        }
    )